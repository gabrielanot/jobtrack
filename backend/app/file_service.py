"""
File Parsing Service
Extract text content from PDF and DOCX files
"""

import io
from typing import Tuple


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_content: Raw bytes of the PDF file
        
    Returns:
        Extracted text content
    """
    from pypdf import PdfReader
    
    reader = PdfReader(io.BytesIO(file_content))
    text_parts = []
    
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_content: Raw bytes of the DOCX file
        
    Returns:
        Extracted text content
    """
    from docx import Document
    
    doc = Document(io.BytesIO(file_content))
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)


def parse_resume_file(filename: str, file_content: bytes) -> Tuple[str, str]:
    """
    Parse a resume file and extract text content.
    
    Args:
        filename: Original filename (used to determine file type)
        file_content: Raw bytes of the file
        
    Returns:
        Tuple of (extracted_text, file_type)
        
    Raises:
        ValueError: If file type is not supported
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        text = extract_text_from_pdf(file_content)
        return text, 'pdf'
    elif filename_lower.endswith('.docx'):
        text = extract_text_from_docx(file_content)
        return text, 'docx'
    elif filename_lower.endswith('.doc'):
        raise ValueError(
            "Old .doc format is not supported. Please convert to .docx or .pdf"
        )
    elif filename_lower.endswith('.txt'):
        text = file_content.decode('utf-8', errors='ignore')
        return text, 'txt'
    else:
        raise ValueError(
            "Unsupported file type. Please upload a PDF, DOCX, or TXT file."
        )
    